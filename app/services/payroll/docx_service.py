"""Payslip generation from an uploaded .docx template.

Two parts:
- ``render_payslip_docx`` fills an uploaded Word template (docxtpl / Jinja2
  tokens) with a payslip context. Pure-Python, always available.
- ``docx_to_pdf`` converts the filled .docx to PDF using LibreOffice headless
  (or, on Windows with MS Word, the ``docx2pdf`` package). This needs a
  converter installed on the host; when none is found it returns ``None`` and
  the caller falls back to the built-in fpdf2 layout.

Token reference for template authors (Jinja2 syntax, e.g. ``{{ net_pay }}``):
    company_name, ref, status, currency, period_start, period_end, pay_date,
    employee.{name,email,code,pan,uan}, company.{name,legal_name,pan,tan},
    gross, total_deductions, net (formatted strings),
    earnings / deductions / employer_contributions: lists of
        {code, label, amount}  -- use a docxtpl table-row loop
    lop_days, paid_days, working_days
"""

from __future__ import annotations

import base64
import io
import os
import re
import shutil
import subprocess  # nosec B404 - used only to invoke the local LibreOffice (soffice) binary
import tempfile
import urllib.request
from html import escape as _escape
from typing import Any
from urllib.parse import urlparse

from docxtpl import DocxTemplate

# OOXML (.docx) files are ZIP archives — they start with the local-file magic.
_ZIP_MAGIC = b"PK\x03\x04"


def looks_like_docx(data: bytes) -> bool:
    return data[:4] == _ZIP_MAGIC


def has_jinja_tokens(data: bytes) -> bool:
    """True if the .docx actually contains ``{{ … }}`` fillable tokens. A
    template with none (e.g. a raw upload that never went through the mapping
    wizard) renders its layout but fills no data — this lets the UI flag that."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return any("{{" in t for t in texts)
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Amount -> words ("Rupees Twelve Thousand … Only")
#
# Many company payslips carry a "Net Pay in words" / "Amount in words" line, so
# we spell the net pay out. Indian-subcontinent currencies use the lakh/crore
# grouping; everyone else uses the international thousand/million grouping. The
# major/minor unit names come per currency code (INR -> Rupees/Paise, etc.).
# ---------------------------------------------------------------------------
_W_ONES = [
    "Zero",
    "One",
    "Two",
    "Three",
    "Four",
    "Five",
    "Six",
    "Seven",
    "Eight",
    "Nine",
    "Ten",
    "Eleven",
    "Twelve",
    "Thirteen",
    "Fourteen",
    "Fifteen",
    "Sixteen",
    "Seventeen",
    "Eighteen",
    "Nineteen",
]
_W_TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

_CURRENCY_UNITS = {
    "INR": ("Rupees", "Paise"),
    "NPR": ("Rupees", "Paisa"),
    "LKR": ("Rupees", "Cents"),
    "PKR": ("Rupees", "Paisa"),
    "BDT": ("Taka", "Poisha"),
    "USD": ("Dollars", "Cents"),
    "AUD": ("Dollars", "Cents"),
    "CAD": ("Dollars", "Cents"),
    "SGD": ("Dollars", "Cents"),
    "HKD": ("Dollars", "Cents"),
    "NZD": ("Dollars", "Cents"),
    "EUR": ("Euros", "Cents"),
    "GBP": ("Pounds", "Pence"),
    "AED": ("Dirhams", "Fils"),
    "SAR": ("Riyals", "Halalas"),
    "QAR": ("Riyals", "Dirhams"),
    "JPY": ("Yen", "Sen"),
    "CNY": ("Yuan", "Fen"),
    "CHF": ("Francs", "Centimes"),
    "ZAR": ("Rand", "Cents"),
    "MYR": ("Ringgit", "Sen"),
    "THB": ("Baht", "Satang"),
    "IDR": ("Rupiah", "Sen"),
    "PHP": ("Pesos", "Centavos"),
}
# Currencies that group digits as thousand / lakh / crore.
_INDIAN_SYSTEM = {"INR", "NPR", "LKR", "PKR", "BDT"}


def _below_thousand(n: int) -> str:
    """0..999 in words ('Three Hundred Forty-Five')."""
    parts: list[str] = []
    if n >= 100:
        parts.append(_W_ONES[n // 100] + " Hundred")
        n %= 100
    if n >= 20:
        word = _W_TENS[n // 10]
        if n % 10:
            word += "-" + _W_ONES[n % 10]
        parts.append(word)
    elif n > 0:
        parts.append(_W_ONES[n])
    return " ".join(parts)


def _int_words_intl(n: int) -> str:
    """Positive int -> words, international (thousand/million/billion) grouping."""
    if n == 0:
        return "Zero"
    scales = ["", " Thousand", " Million", " Billion", " Trillion", " Quadrillion"]
    chunks: list[str] = []
    scale = 0
    while n > 0 and scale < len(scales):
        chunk = n % 1000
        if chunk:
            chunks.append(_below_thousand(chunk) + scales[scale])
        n //= 1000
        scale += 1
    return " ".join(reversed(chunks))


def _int_words_indian(n: int) -> str:
    """Positive int -> words, Indian (thousand/lakh/crore) grouping."""
    if n == 0:
        return "Zero"
    parts: list[str] = []
    crore, n = divmod(n, 10**7)
    lakh, n = divmod(n, 10**5)
    thousand, rest = divmod(n, 1000)
    if crore:
        parts.append(_int_words_indian(crore) + " Crore")  # handles 100+ crore
    if lakh:
        parts.append(_below_thousand(lakh) + " Lakh")
    if thousand:
        parts.append(_below_thousand(thousand) + " Thousand")
    if rest:
        parts.append(_below_thousand(rest))
    return " ".join(parts)


def amount_to_words(amount: float | int | str | None, currency: str = "INR") -> str:
    """Spell out a money amount, e.g. ``12345.50`` / ``INR`` ->
    "Rupees Twelve Thousand Three Hundred Forty-Five and Fifty Paise Only".
    Unknown currency codes fall back to the code itself + "Cents"."""
    try:
        value = float(amount or 0)
    except (TypeError, ValueError):
        return ""
    code = (currency or "INR").strip().upper()
    major_unit, minor_unit = _CURRENCY_UNITS.get(code, (code or "Units", "Cents"))
    negative = value < 0
    value = abs(value)
    whole = int(value)
    minor = int(round((value - whole) * 100))
    if minor >= 100:  # rounding carried into the next whole unit
        whole += 1
        minor -= 100
    int_words = _int_words_indian(whole) if code in _INDIAN_SYSTEM else _int_words_intl(whole)
    text = f"{major_unit} {int_words}"
    if minor:
        text += f" and {_below_thousand(minor)} {minor_unit}"
    text += " Only"
    return ("Minus " + text) if negative else text


def render_payslip_docx(
    template_bytes: bytes, context: dict[str, Any], logo_image: bytes | None = None
) -> bytes:
    """Fill the uploaded .docx template with ``context`` and return .docx bytes.

    When ``logo_image`` (PNG bytes) is given, the company logo is embedded **at
    the template's logo placeholder** — a ``[ COMPANY LOGO ]`` line, a bare
    "Company Logo" line, or a ``{{ logo }}`` token — anywhere in the body, a
    table, or a section header/footer. The placeholder text is removed and the
    picture takes its place (keeping the placeholder paragraph's alignment), so
    the logo never lands in a stray spot. If no placeholder is found we fall back
    to the ``{{ logo }}`` token via docxtpl so existing mappings still work."""
    doc = DocxTemplate(io.BytesIO(template_bytes))
    ctx = dict(context)
    if logo_image:
        placed = False
        try:
            placed = _place_logo_image(doc.get_docx(), logo_image)
        except Exception:
            placed = False
        if placed:
            ctx["logo"] = ""  # the real image already replaced the placeholder
        else:
            try:
                from docx.shared import Mm
                from docxtpl import InlineImage

                ctx["logo"] = InlineImage(doc, io.BytesIO(logo_image), width=Mm(38))
            except Exception:
                pass  # fall back to whatever text 'logo' the context already holds
    doc.render(ctx)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# Short standalone phrases (normalised) that mark where a logo belongs.
_LOGO_PHRASES = {
    "logo",
    "company logo",
    "logo here",
    "company logo here",
    "your logo",
    "company logo name",
    "logo image",
    "company name logo",
    "insert logo",
    "insert company logo",
    "logo goes here",
}


def _logo_placeholder_priority(text: str) -> int:
    """How strongly a paragraph reads as the logo placeholder (0 = not one).
    Higher wins: a bracketed ``[ COMPANY LOGO ]`` is the clearest designed spot,
    a bare logo phrase next, a stray ``{{ logo }}`` token last (it may be a
    mis-mapped slot)."""
    t = (text or "").strip()
    if not t:
        return 0
    if t[0] == "[" and t[-1] == "]" and "logo" in t.lower():
        return 3
    if _normalize(t) in _LOGO_PHRASES:
        return 2
    if re.search(r"\{\{\s*logo\s*\}\}", t):
        return 1
    return 0


def _iter_logo_paragraphs(document):
    """Every paragraph anywhere the logo could be: body + tables (nested too),
    AND each section's headers/footers — which docxtpl and the mapping wizard
    never reach. Uses native python-docx containers so an inserted picture
    attaches to the right document part (body vs. header/footer)."""

    def _walk(container):
        for p in container.paragraphs:
            yield p
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _walk(cell)

    yield from _walk(document)
    seen: set[int] = set()
    for section in document.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if hf is None:
                continue
            key = id(hf._element)
            if key in seen:
                continue
            seen.add(key)
            try:
                yield from _walk(hf)
            except Exception:
                continue  # nosec B112 - best-effort recursion; skip unreadable nodes


def _place_logo_image(document, logo_image: bytes) -> bool:
    """Embed ``logo_image`` at the template's logo placeholder. The strongest
    placeholder gets the picture; every other logo placeholder/token is cleared
    so the logo never appears twice or in a stray spot. Returns True if placed."""
    from docx.shared import Mm

    candidates: list[tuple[int, Any]] = []
    for para in _iter_logo_paragraphs(document):
        pr = _logo_placeholder_priority(para.text)
        if pr:
            candidates.append((pr, para))
    if not candidates:
        return False
    candidates.sort(key=lambda c: c[0], reverse=True)
    target = candidates[0][1]
    for _, para in candidates:  # wipe brackets/tokens/duplicates first
        _set_para_text(para, "")
    target.add_run().add_picture(io.BytesIO(logo_image), width=Mm(38))
    return True


# Logos rarely change; cache the normalised bytes per URL for the process.
_LOGO_CACHE: dict[str, bytes | None] = {}


def fetch_logo_image(url: str | None) -> bytes | None:
    """Fetch the company logo from its URL and return normalised PNG bytes
    (resized small), or None if unset/unreachable/not an image. Only http(s)/data
    URLs are honoured; all errors are swallowed so rendering never fails on it."""
    if not url:
        return None
    if url in _LOGO_CACHE:
        return _LOGO_CACHE[url]
    data: bytes | None = None
    try:
        if urlparse(url).scheme in ("http", "https", "data"):
            req = urllib.request.Request(url, headers={"User-Agent": "payroll"})
            with urllib.request.urlopen(req, timeout=6) as resp:  # nosec B310 - scheme allow-listed (http/https/data) above
                raw = resp.read(4 * 1024 * 1024 + 1)
            if 0 < len(raw) <= 4 * 1024 * 1024:
                from PIL import Image

                im = Image.open(io.BytesIO(raw))
                im.load()
                if im.mode not in ("RGB", "RGBA"):
                    im = im.convert("RGBA")
                im.thumbnail((160, 70))  # keep it logo-sized
                buf = io.BytesIO()
                im.save(buf, "PNG")
                data = buf.getvalue()
    except Exception:
        data = None
    _LOGO_CACHE[url] = data
    return data


def sample_payslip_template() -> bytes:
    """A ready-to-use .docx payslip template with correct tokens.

    Built with python-docx so every token sits in a single run (no Word
    run-splitting), which is the usual reason hand-typed tokens don't fill.
    Admins download this, restyle it, and re-upload. Uses the preformatted
    ``*_lines`` tokens (reliable) rather than docxtpl's finicky ``{%tr%}`` loops.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = Document()
    title = d.add_paragraph()
    run = title.add_run("{{ company_name }}")
    run.bold = True
    run.font.size = Pt(18)
    sub = d.add_paragraph("PAYSLIP — Ref #{{ ref }}  ·  Status: {{ status }}")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    d.add_paragraph("Employee: {{ employee.name }}  ({{ employee.code }})")
    d.add_paragraph("Email: {{ employee.email }}")
    d.add_paragraph("Pay period: {{ period_start }} to {{ period_end }}   ·   Pay date: {{ pay_date }}")

    d.add_paragraph("")
    h1 = d.add_paragraph().add_run("Earnings")
    h1.bold = True
    d.add_paragraph("{{ earnings_lines }}")
    d.add_paragraph("Gross Earnings\t{{ gross }}")

    d.add_paragraph("")
    h2 = d.add_paragraph().add_run("Deductions")
    h2.bold = True
    d.add_paragraph("{{ deductions_lines }}")
    d.add_paragraph("Total Deductions\t{{ total_deductions }}")

    d.add_paragraph("")
    h3 = d.add_paragraph().add_run("Attendance")
    h3.bold = True
    d.add_paragraph("Working days: {{ working_days }}   LOP: {{ lop_days }}   Paid: {{ paid_days }}")

    d.add_paragraph("")
    net = d.add_paragraph()
    nrun = net.add_run("NET PAYABLE\t{{ net }}")
    nrun.bold = True
    nrun.font.size = Pt(14)

    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Smart mapping wizard
#
# Companies usually have an existing payslip Word document — styled with their
# logo and a table of labels ("Basic", "Net Pay", "Employee Name") but with
# blank/sample value cells and NO docxtpl tokens. The wizard bridges that:
#   1. ``scan_docx_fields`` detects the value "slots" (table value-cells and
#      "Label: …" paragraphs) and suggests which payroll field each maps to.
#   2. The admin reviews/edits the mapping in the UI.
#   3. ``apply_field_mapping`` injects the chosen ``{{ token }}`` into each slot
#      (single-run, so docxtpl reliably fills it) and returns a tokenised .docx
#      that the normal generation path renders.
# The traversal in ``_iter_doc_slots`` is deterministic, so scan and apply see
# the same ordered slots — the UI only needs to send {slot_index: token}.
# ---------------------------------------------------------------------------

# Catalogue of mappable payroll fields: (group, token key, human label, synonyms).
# ``key`` is the docxtpl token (see ``_docx_context``). ``synonyms`` drive
# auto-detection of a label cell; they must stay unique across fields.
_FIELDS: list[tuple[str, str, str, list[str]]] = [
    # --- Employee ---
    (
        "Employee",
        "employee.name",
        "Employee name",
        ["employee name", "name of employee", "name", "emp name", "staff name"],
    ),
    (
        "Employee",
        "employee.code",
        "Employee code / ID",
        [
            "employee code",
            "employee id",
            "emp id",
            "emp code",
            "employee no",
            "staff id",
            "emp no",
            "employee number",
        ],
    ),
    ("Employee", "employee.designation", "Designation", ["designation", "title", "job title", "role"]),
    ("Employee", "employee.department", "Department", ["department", "dept"]),
    ("Employee", "employee.location", "Location / Branch", ["location", "work location", "branch", "office"]),
    ("Employee", "employee.email", "Email", ["email", "email id", "e mail"]),
    (
        "Employee",
        "employee.bank_account_no",
        "Bank account no.",
        [
            "bank account",
            "account no",
            "account number",
            "a c no",
            "bank a c no",
            "bank account number",
            "bank a c",
        ],
    ),
    ("Employee", "employee.pan", "Employee PAN", ["pan", "pan no", "pan number"]),
    ("Employee", "employee.uan", "UAN", ["uan", "uan no", "uan number"]),
    ("Employee", "employee.esic", "ESIC number", ["esic", "esic no", "esi number", "esic number"]),
    ("Employee", "employee.state", "Employee state", ["state"]),
    ("Employee", "employee.date_of_joining", "Date of joining", ["date of joining", "doj", "joining date"]),
    # --- Pay period ---
    (
        "Pay period",
        "cycle_name",
        "Pay period / month",
        ["pay period", "period", "month", "salary month", "for the month", "pay month", "salary period"],
    ),
    ("Pay period", "period_start", "Period start", ["period start", "from date", "start date"]),
    ("Pay period", "period_end", "Period end", ["period end", "to date", "end date"]),
    ("Pay period", "pay_date", "Pay date", ["pay date", "payment date", "date of payment", "paid on"]),
    ("Pay period", "ref", "Reference no.", ["ref", "reference", "payslip no", "reference no", "slip no"]),
    # --- Attendance ---
    ("Attendance", "working_days", "Working days", ["working days", "total days", "days in month"]),
    ("Attendance", "paid_days", "Paid days", ["paid days", "days paid", "days worked"]),
    ("Attendance", "lop_days", "Loss-of-pay days", ["lop", "lop days", "loss of pay", "absent days"]),
    # --- Earnings (component amounts) ---
    ("Earnings", "amount.BASIC", "Basic", ["basic", "basic salary", "basic pay", "basic wage"]),
    ("Earnings", "amount.HRA", "HRA", ["hra", "house rent allowance", "house rent"]),
    ("Earnings", "amount.DA", "Dearness allowance", ["da", "dearness allowance"]),
    (
        "Earnings",
        "amount.CONVEYANCE",
        "Conveyance",
        ["conveyance", "conveyance allowance", "transport allowance", "transport"],
    ),
    ("Earnings", "amount.MEDICAL", "Medical allowance", ["medical", "medical allowance"]),
    ("Earnings", "amount.LTA", "LTA", ["lta", "leave travel allowance"]),
    (
        "Earnings",
        "amount.SPECIAL",
        "Special allowance",
        ["special allowance", "special", "other allowance", "spl allowance"],
    ),
    # --- Deductions (component amounts) ---
    (
        "Deductions",
        "amount.PF",
        "Provident fund (PF)",
        ["pf", "provident fund", "epf", "employee pf", "pf contribution"],
    ),
    ("Deductions", "amount.ESI", "ESI", ["employee esi", "esi contribution"]),
    ("Deductions", "amount.PT", "Professional tax", ["professional tax", "pt", "p tax", "prof tax"]),
    ("Deductions", "amount.TDS", "Income tax (TDS)", ["tds", "income tax", "tax deducted", "i tax"]),
    # --- Totals ---
    (
        "Totals",
        "gross",
        "Gross earnings",
        ["gross", "gross earnings", "gross salary", "total earnings", "total gross"],
    ),
    ("Totals", "total_deductions", "Total deductions", ["total deductions", "deductions", "total deduction"]),
    (
        "Totals",
        "net",
        "Net pay",
        ["net pay", "net salary", "net payable", "take home", "amount payable", "net amount", "net"],
    ),
    (
        "Totals",
        "net_in_words",
        "Net pay (in words)",
        [
            "net pay in words",
            "amount in words",
            "net amount in words",
            "in words",
            "rupees in words",
            "net salary in words",
            "net payable in words",
            "amount payable in words",
            "amount in word",
            "rupees",
            "in words only",
        ],
    ),
    # --- Lists (multi-line blocks) ---
    ("Lists", "earnings_lines", "All earnings (multi-line block)", []),
    ("Lists", "deductions_lines", "All deductions (multi-line block)", []),
    ("Lists", "employer_contributions_lines", "Employer contributions (multi-line block)", []),
    # --- Company ---
    ("Company", "company_name", "Company name", ["company name", "organisation", "organization", "company"]),
    (
        "Company",
        "logo",
        "Company logo / name",
        ["logo", "company logo", "logo here", "your logo", "company logo here"],
    ),
    ("Company", "company.address", "Company address", ["company address", "address", "registered address"]),
    ("Company", "company.pan", "Company PAN", ["company pan"]),
    ("Company", "company.tan", "Company TAN", ["tan", "tan no", "company tan"]),
    ("Misc", "currency", "Currency", ["currency"]),
    ("Misc", "status", "Status", ["status"]),
]

# Public catalogue (no synonyms) for the UI's field dropdown.
FIELD_CATALOG: list[dict[str, str]] = [{"group": g, "key": k, "label": label} for g, k, label, _ in _FIELDS]


def _normalize(s: str) -> str:
    """Loosely normalise a label for synonym matching: lowercase, drop
    parentheticals/punctuation, collapse whitespace."""
    s = s.lower()
    s = re.sub(r"\(.*?\)", " ", s)  # drop "(₹)" / "(per month)" etc.
    s = re.sub(r"[^a-z0-9 ]", " ", s)  # keep alphanumerics + spaces
    s = re.sub(r"\s+", " ", s).strip()
    return s


# normalized synonym -> token key. Built once from _FIELDS.
_SYNONYMS: dict[str, str] = {}
for _g, _k, _label, _syns in _FIELDS:
    for _syn in _syns:
        _SYNONYMS.setdefault(_normalize(_syn), _k)


def _set_para_text(paragraph, text: str) -> None:
    """Replace a paragraph's text with a single run (preserving the first run's
    formatting). Single-run avoids Word's token-splitting across runs."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paragraph.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    _set_para_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        _set_para_text(extra, "")


def _row_text(cells) -> str:
    """A readable 'A | B | C' summary of a row (de-duping merged cells)."""
    out: list[str] = []
    prev = None
    for c in cells:
        t = c.text.strip()
        if t and t != prev:
            out.append(t)
        prev = t
    return " | ".join(out)


# A cell counts as "blank" if it's empty or just placeholder filler that real
# company templates leave for the value (underscores, dashes, dotted lines, 0).
def _is_blankish(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    if re.fullmatch(r"[_.\-–—\s·•]*", t):
        return True
    return t in {"0", "0.0", "0.00", "-", "--", "nil", "n/a", "na", "xxxx"}


# A cell is "value-like" (the place an amount/value goes) if it's blank or looks
# like a number/amount — e.g. "5,000.00", "₹ 5000", "Rs. 1200" (few letters).
def _is_value_like(text: str) -> bool:
    t = text.strip()
    if _is_blankish(t):
        return True
    if not any(c.isdigit() for c in t):
        return False
    return sum(1 for c in t if c.isalpha()) <= 3  # tolerate "Rs"/"INR"/currency


# A cell is "label-like" if it has real words and isn't itself a value.
def _is_label_like(text: str) -> bool:
    t = text.strip()
    if not t or _is_value_like(t):
        return False
    return any(c.isalpha() for c in t)


# Detect a Jinja token already present in a value, e.g. "{{ amount.BASIC }}", so
# a template that's already (partly) tokenised pre-maps itself in the wizard.
_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\}\}")


def _existing_token(text: str | None) -> str | None:
    m = _TOKEN_RE.search(text or "")
    return m.group(1) if m else None


def _distinct_cells(row):
    """Row cells with consecutive horizontally-merged duplicates collapsed, so a
    merged cell counts as one visual column."""
    out = []
    for c in row.cells:
        if out and out[-1] is c:
            continue
        out.append(c)
    return out


def _make_table_slot(idx: int, label_cell, target_cell, row_cells, token):
    """Build a (slot, apply_fn) tuple for a table label→value pair."""
    return (
        {
            "index": idx,
            "kind": "table",
            "label": label_cell.text.strip(),
            "context": _row_text(row_cells),
            "current": target_cell.text.strip(),
            "suggested_token": token,
        },
        (lambda c: lambda tok: _set_cell_text(c, "{{ " + tok + " }}"))(target_cell),
    )


def _paragraph_slot(paragraph, suggest):
    """Return ``(slot_without_index, apply_fn)`` for a paragraph that is a
    bracketed placeholder or a ``Label: value`` line, else ``None``. Works for
    body paragraphs AND paragraphs inside a table cell."""
    text = paragraph.text
    stripped = text.strip()
    # Bracketed placeholder like "[ COMPANY LOGO ]".
    if len(stripped) > 2 and stripped[0] == "[" and stripped[-1] == "]":
        inner = stripped[1:-1].strip()
        token = _existing_token(stripped) or suggest(inner)
        return (
            {
                "kind": "placeholder",
                "label": inner or stripped,
                "context": stripped,
                "current": stripped,
                "suggested_token": token,
            },
            (lambda par: lambda tok: _set_para_text(par, "{{ " + tok + " }}"))(paragraph),
        )
    # "Label: value" — blank or short value (skip long sentences).
    if ":" in text:
        label, _, rest = text.partition(":")
        if _is_label_like(label) and (_is_blankish(rest) or len(rest.split()) <= 4):
            token = _existing_token(rest) or suggest(label)
            return (
                {
                    "kind": "para",
                    "label": label.strip(),
                    "context": text.strip(),
                    "current": rest.strip(),
                    "suggested_token": token,
                },
                (lambda par, lbl: lambda tok: _set_para_text(par, lbl + ": {{ " + tok + " }}"))(
                    paragraph, label.strip()
                ),
            )
    return None


def _iter_doc_slots(doc, synonyms: dict[str, str] | None = None):
    """Yield ``(slot, apply_fn)`` in a deterministic order.

    A *slot* is a place a value belongs: a table value-cell next to a label, or
    a ``Label: …`` paragraph. We surface **every** label→blank pair (not only
    ones we recognise) so a company's own labels — "Performance Bonus", "Special
    City Allowance" — still appear for the admin to map; ``suggested_token`` is
    pre-filled when we recognise the label, else null. ``apply_fn(token)`` writes
    ``{{ token }}`` there.

    Crucially, **which** slots are emitted depends only on the document's
    structure (a label beside a blank/numeric value), NOT on ``synonyms`` — so a
    slot's index is identical at scan time and apply time even if the recognised
    synonyms differ. ``synonyms`` only populates the (cosmetic) suggestion."""
    idx = 0
    synonyms = synonyms or {}

    def suggest(text: str) -> str | None:
        return synonyms.get(_normalize(text))

    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    def _has_colon(cell) -> bool:
        return ":" in cell.text

    def _process_table(table):
        """Column-pair COLON-LESS label→value cells (the amounts/details grid).
        Cells that contain a colon ("Label: value") are left for the paragraph
        pass below, so the two passes never emit the same field twice. Recurses
        into nested tables."""
        nonlocal idx
        rows = list(table.rows)
        consumed: set[int] = set()  # rows used as the "values" of a stacked pair
        for r, row in enumerate(rows):
            if r in consumed:
                continue
            cells = _distinct_cells(row)
            # (b) Vertical: an all-label row directly above an all-blank row.
            if r + 1 < len(rows):
                below = _distinct_cells(rows[r + 1])
                if (
                    len(cells) >= 2
                    and len(below) == len(cells)
                    and all(_is_label_like(c.text) and not _has_colon(c) for c in cells)
                    and all(_is_blankish(c.text) for c in below)
                ):
                    for lbl, tgt in zip(cells, below):
                        token = _existing_token(tgt.text) or suggest(lbl.text)
                        yield _make_table_slot(idx, lbl, tgt, cells, token)
                        idx += 1
                    consumed.add(r + 1)
                    continue
            # (a) Horizontal: pair columns (0,1), (2,3), … — only colon-less cells.
            i = 0
            while i + 1 < len(cells):
                lbl, tgt = cells[i], cells[i + 1]
                if _is_label_like(lbl.text) and not _has_colon(lbl) and not _has_colon(tgt):
                    token = _existing_token(tgt.text) or suggest(lbl.text)
                    yield _make_table_slot(idx, lbl, tgt, cells, token)
                    idx += 1
                i += 2
            # Recurse into nested tables.
            for cell in cells:
                for nested in cell.tables:
                    yield from _process_table(nested)

    # 1) Tables — colon-less label→value grids (amounts, stacked details).
    for table in doc.tables:
        yield from _process_table(table)

    # 2) EVERY paragraph in the document (body, table cells at any depth, text
    #    boxes) that is a "Label: value" line or a "[ placeholder ]". Iterating
    #    the XML catches paragraphs that doc.paragraphs / doc.tables miss. These
    #    are disjoint from pass 1, which only paired colon-less cells.
    for p_elem in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p_elem, doc)
        res = _paragraph_slot(para, suggest)
        if res:
            slot, fn = res
            slot["index"] = idx
            yield (slot, fn)
            idx += 1


# Cap detected slots so a pathological document can't blow up the payload.
_MAX_SLOTS = 300


def scan_docx_fields(data: bytes, extra_synonyms: dict[str, str] | None = None) -> list[dict[str, Any]]:
    """Detect value slots in an uploaded .docx and suggest a field for each.

    ``extra_synonyms`` (normalized label -> token) lets the caller add the
    company's own salary-component labels so they auto-suggest too."""
    from docx import Document

    synonyms = dict(_SYNONYMS)
    if extra_synonyms:
        synonyms.update(extra_synonyms)
    doc = Document(io.BytesIO(data))
    return [slot for slot, _ in _iter_doc_slots(doc, synonyms)][:_MAX_SLOTS]


def apply_field_mapping(data: bytes, mapping: dict[int, str]) -> bytes:
    """Inject ``{{ token }}`` into each mapped slot and return tokenised .docx.

    ``mapping`` maps a slot index (from ``scan_docx_fields``) to a token key;
    unmapped/blank slots are left untouched."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    for slot, apply in _iter_doc_slots(doc):
        token = mapping.get(slot["index"])
        if token:
            apply(token)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# docx -> HTML
#
# So the company's template can drive BOTH the on-screen / print payslip view
# AND the PDF (via fpdf2's HTML engine) even when no LibreOffice/Word converter
# is installed. The HTML is intentionally simple: inline ``style`` for the
# browser, plus the attributes fpdf2 understands (``border``/``width`` on
# tables). Bold/italic/underline and paragraph alignment are preserved.
# ---------------------------------------------------------------------------
def _run_image_html(run) -> str:
    """Emit any inline image in this run as a self-contained <img> (data URI), so
    an embedded logo shows in the web/print view and the fpdf2 PDF."""
    try:
        from docx.oxml.ns import qn
    except Exception:
        return ""
    blips = run._element.findall(".//" + qn("a:blip"))
    if not blips:
        return ""
    out = []
    for blip in blips:
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid:
            continue
        try:
            part = run.part.related_parts[rid]
            blob = part.blob
            ctype = getattr(part, "content_type", None) or "image/png"
            dim = ""
            try:
                from PIL import Image

                w, h = Image.open(io.BytesIO(blob)).size
                dim = f' width="{w}" height="{h}"'
            except Exception:
                pass
            b64 = base64.b64encode(blob).decode("ascii")
            out.append(f'<img src="data:{ctype};base64,{b64}"{dim} style="max-height:70px;"/>')
        except Exception:
            pass
    return "".join(out)


def _run_html(run) -> str:
    img = _run_image_html(run)
    t = _escape(run.text).replace("\n", "<br>").replace("\t", "&emsp;")
    if t:
        if run.bold:
            t = f"<b>{t}</b>"
        if run.italic:
            t = f"<i>{t}</i>"
        if run.underline:
            t = f"<u>{t}</u>"
    return img + t


# python-docx alignment enum -> CSS. (LEFT=0, CENTER=1, RIGHT=2, JUSTIFY=3)
_ALIGN_CSS = {1: "center", 2: "right", 3: "justify"}


def _para_html(p) -> str:
    body = "".join(_run_html(r) for r in p.runs)
    style = "margin:4px 0;"
    if p.alignment is not None:
        css = _ALIGN_CSS.get(int(p.alignment))
        if css:
            style += f"text-align:{css};"
    if not body.strip():
        return f'<p style="{style}">&nbsp;</p>'
    return f'<p style="{style}">{body}</p>'


def _table_html(table) -> str:
    rows_html = []
    for row in table.rows:
        cells_html = []
        seen: set[int] = set()  # de-dupe horizontally-merged cells
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            inner = "".join(_para_html(p) for p in cell.paragraphs) or "&nbsp;"
            cells_html.append(f'<td style="border:1px solid #d1d5db;padding:4px 8px;">{inner}</td>')
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return (
        '<table border="1" width="100%" '
        'style="width:100%;border-collapse:collapse;margin:6px 0;">' + "".join(rows_html) + "</table>"
    )


def docx_to_html(docx_bytes: bytes) -> str:
    """Render a (filled) .docx to a simple HTML fragment for screen/print/PDF."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    # Iterate the body in document order so paragraphs and tables interleave
    # exactly as authored.
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            parts.append(_para_html(Paragraph(child, doc)))
        elif tag == "tbl":
            parts.append(_table_html(Table(child, doc)))
    return '<div class="payslip-doc">' + "".join(parts) + "</div>"


def _find_soffice() -> str | None:
    """Locate LibreOffice: explicit override, then PATH, then common installs."""
    override = os.getenv("PAYROLL_SOFFICE_PATH")
    if override and os.path.exists(override):
        return override
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for guess in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
    ):
        if os.path.exists(guess):
            return guess
    return None


def pdf_conversion_available() -> bool:
    """True if a docx->pdf converter (LibreOffice or MS Word) is on this host."""
    if _find_soffice():
        return True
    try:  # MS Word via docx2pdf (Windows/macOS, optional dependency)
        import docx2pdf  # noqa: F401

        return True
    except Exception:
        return False


def docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Convert .docx bytes to PDF. Returns None when conversion is unavailable
    or fails, so the caller can fall back to the built-in PDF layout."""
    soffice = _find_soffice()
    if soffice:
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "payslip.docx")
            with open(src, "wb") as fh:
                fh.write(docx_bytes)
            # A per-call user profile avoids lock clashes under concurrency.
            profile = os.path.join(tmp, "profile")
            try:
                subprocess.run(  # nosec B603 - fixed soffice binary + controlled args, no shell
                    [
                        soffice,
                        f"-env:UserInstallation=file:///{profile.replace(os.sep, '/')}",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        tmp,
                        src,
                    ],
                    check=True,
                    capture_output=True,
                    timeout=90,
                )
            except (subprocess.SubprocessError, OSError):
                return None
            pdf_path = os.path.join(tmp, "payslip.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as fh:
                    return fh.read()
        return None

    # Fallback: MS Word via docx2pdf, if installed.
    try:
        import docx2pdf
    except Exception:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "payslip.docx")
        dst = os.path.join(tmp, "payslip.pdf")
        with open(src, "wb") as fh:
            fh.write(docx_bytes)
        try:
            docx2pdf.convert(src, dst)
        except Exception:
            return None
        if os.path.exists(dst):
            with open(dst, "rb") as fh:
                return fh.read()
    return None
